import markdown
from docx import Document
import os

def markdown_to_word(markdown_file_path, output_filename):
    # 从文件中读取 Markdown 内容
    with open(markdown_file_path, 'r', encoding='utf-8') as file:
        markdown_text = file.read()

    # 将 Markdown 转换为 HTML
    html = markdown.markdown(markdown_text)

    # 创建一个新的 Word 文档
    doc = Document()

    # 简单解析 HTML 并添加到 Word 文档
    # 注意：这个解析是基础的，可能不支持复杂的 HTML 结构
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')

    for element in soup.find_all():
        if element.name == 'p':
            # 处理段落
            doc.add_paragraph(element.get_text(strip=True))
        elif element.name == 'h1':
            # 处理标题 1
            doc.add_heading(element.get_text(strip=True), level=1)
        elif element.name == 'h2':
            # 处理标题 2
            doc.add_heading(element.get_text(strip=True), level=2)
        # 可以继续添加其他 HTML 标签的处理

    # 保存 Word 文档
    doc.save(output_filename)
    print(f"Word 文件已保存：{output_filename}")

# Markdown 文件路径
markdown_file_path = '/Users/caijunxiang/IdeaProjects/ylz_code/data-asset/data-asset-application/data-asset-admin/data-asset-admin-api/src/main/java/com/yhtech/dam/admin/api/service/readme.md'  # 确保这个路径指向你的 Markdown 文件

# 输出的 Word 文件名
output_filename = '/Users/caijunxiang/IdeaProjects/ylz_code/data-asset/data-asset-application/data-asset-admin/data-asset-admin-api/src/main/java/com/yhtech/dam/admin/api/service/readme.docs'

# 执行转换
markdown_to_word(markdown_file_path, output_filename)

