from docx import Document
import uuid
import datetime

def word_tables_to_txt(word_path, txt_path):
    """
    从Word文档中提取所有表格内容，并写入TXT文件
    :param word_path: Word文件路径（.docx格式）
    :param txt_path: 输出的TXT文件路径
    """
    # 读取Word文档
    doc = Document(word_path)

    with open(txt_path, 'w', encoding='utf-8') as f:
        # 遍历所有表格
        for table in doc.tables:
            for row in table.rows:
                # 获取每行的单元格内容
                row_text = [cell.text.strip() for cell in row.cells]
                # 用制表符分隔单元格内容并写入文件
                f.write('\t'.join(row_text) + '\n')
            # 表格之间用分隔线区分
            f.write('-' * 50 + '\n')

    print(f"表格内容已成功导出到 {txt_path}")


def word_tables_to_sql(word_path, txt_path, table_name='das_biz.das_catalogue_ds_item'):
    """
    从Word文档中提取表格内容并生成符合das_catalogue_ds_item表结构的SQL INSERT语句
    :param word_path: Word文件路径(.docx格式)
    :param txt_path: 输出的SQL文件路径
    :param table_name: 生成的表名，默认为'das_biz.das_catalogue_ds_item'
    """
    doc = Document(word_path)

    # 数据类型映射字典
    type_mapping = {
        '字符串': '15',
        '数值': '4',
        '日期': '93',
        '时间': '92',
        '日期时间': '93',
        '布尔': '16'
    }

    # 固定参数
    fixed_params = {
        'create_by': '1947902981197328386',
        'del_flag': '0',
        'item_precision': '0',
        'item_dict': '',
        'update_by': 'NULL',
        'update_time': 'NULL'
    }

    with open(txt_path, 'w', encoding='utf-8') as f:
        for table_idx, table in enumerate(doc.tables):
            # 写入注释说明
            f.write(f"-- 表{table_idx+1}的INSERT语句\n")
            f.write(f"-- 表名: {table_name}\n\n")

            # 处理数据行(从第二行开始，假设第一行是表头)
            for row in table.rows[1:]:
                # 获取单元格数据
                cells = [cell.text.strip() for cell in row.cells]

                if len(cells) < 6:  # 确保至少有6列数据
                    continue

                # 解析表格数据
                data_item = {
                    'item_name': cells[0],  # 数据项
                    'item_en_name': cells[1],  # 字段名
                    'item_type': type_mapping.get(cells[2], '15'),  # 数据类型
                    'item_length': cells[3],  # 类型长度
                    'item_nullable': '0' if cells[4] == '必填' else '1',  # 填报要求
                    'item_key_flag': '1' if '主键' in cells[5] else '0',  # 是否主键
                    'create_time': datetime.datetime.now().strftime('%Y%m%d%H%M%S')
                }

                # 生成ID
                data_item['id'] = str(uuid.uuid4().int)[:19]
                # 假设das_catalogue_ds_id是固定的或需要从其他地方获取
                data_item['das_catalogue_ds_id'] = '1947933443090419712'

                # 合并固定参数
                data_item.update(fixed_params)

                # 生成INSERT语句
                columns = [
                    'id', 'das_catalogue_ds_id', 'item_name', 'item_en_name',
                    'item_type', 'item_length', 'item_precision', 'item_nullable',
                    'item_dict', 'create_by', 'create_time', 'update_by',
                    'update_time', 'del_flag', 'item_key_flag'
                ]

                values = []
                for col in columns:
                    val = data_item.get(col, 'NULL')
                    if val != 'NULL' and not (isinstance(val, str) and val.isdigit()):
                        val = f"'{val}'"
                    values.append(str(val))

                sql = f"INSERT INTO {table_name} ({', '.join(columns)}) VALUES ({', '.join(values)});\n"
                f.write(sql)

            # 表格之间用空行分隔
            f.write('\n')

    print(f"SQL语句已成功导出到 {txt_path}")



